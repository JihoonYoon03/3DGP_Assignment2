# -*- coding: utf-8 -*-
"""3D게임 프로그래밍 과제 2 설명 문서 .docx 생성기.

1차 보고서(첨부 PDF)의 레이아웃을 그대로 따르되, 구현 내용은 코드 순서 나열 대신
"무엇을 어떤 순서로 구현했는지" 의 서술형으로 구성. 코드는 핵심 기법 전달이 필요할 때만 첨부.
"""

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Users\wlgns\Documents\Github\3DGP_Assignment2\.claude\worktrees\cranky-murdock-ab8e03\_report\2차과제보고서.docx"

# ----- 폰트 상수 -----
FONT_BODY = "맑은 고딕"        # 본문/제목
FONT_CODE = "굴림체"           # 코드 블록 (한글 모노스페이스)

# ----- 문서 생성 -----
doc = Document()

# A4 + 1차 보고서와 비슷한 여백
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# 기본 스타일을 맑은 고딕 11pt 로
style = doc.styles["Normal"]
style.font.name = FONT_BODY
style.font.size = Pt(11)
# 동아시아 폰트도 명시
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), FONT_BODY)
rfonts.set(qn("w:ascii"), FONT_BODY)
rfonts.set(qn("w:hAnsi"), FONT_BODY)


def set_korean_font(run, font_name=FONT_BODY):
    """run 에 한글 동아시아 폰트를 명시한다."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_para(text, *, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
             first_line_indent=None, space_before=0, space_after=4,
             font=FONT_BODY, line_spacing=1.4):
    p = doc.add_paragraph()
    p.alignment = align
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        set_korean_font(r, font)
    return p


def add_section_heading(text):
    """본문 섹션 제목 (예: "1. 과제 목표")."""
    add_para(text, bold=True, size=14, space_before=14, space_after=8)


def add_subsection_heading(text):
    """하위 항목 제목 (예: "1) 씬 구조 분리와 흐름 제어")."""
    add_para(text, bold=True, size=12, space_before=10, space_after=5)


def add_body(text):
    """본문 단락. 첫 줄 들여쓰기 약간."""
    add_para(text, size=11, first_line_indent=Pt(11), space_after=5)


def add_bullet(text):
    """들여쓰기된 항목 라인."""
    add_para("· " + text, size=11, space_after=2)


def add_code(text):
    """코드 블록. 굴림체(한글 모노스페이스) + 작은 글자 + 회색 배경."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    # 회색 배경
    ppr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    ppr.append(shd)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = FONT_CODE
        r.font.size = Pt(8.5)
        set_korean_font(r, FONT_CODE)
        if i < len(lines) - 1:
            r.add_break()


def add_page_break():
    doc.add_page_break()


# ====================================================================
# ---------- 표지 ----------
# ====================================================================
# 1차 보고서와 같은 레이아웃 — 페이지 중상단에 제목, 그 아래 몇 줄 띄우고 이름.
# A4 기준 페이지 한 장에 모두 들어가도록 빈 줄 수를 신중히 조정.
for _ in range(9):
    add_para("", size=11)
add_para("3D게임 프로그래밍 과제 2 설명 문서",
         bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
for _ in range(7):
    add_para("", size=11)
add_para("2022184025 윤지훈",
         size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

add_page_break()

# ====================================================================
# ---------- 목차 ----------
# ====================================================================
add_para("", size=11)
add_para("목차", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

for item in ["1. 과제 목표", "2. 가정", "3. 조작법", "4. 실행 결과", "5. 구현 내용", "6. 미흡한 점"]:
    add_para(item, size=14, space_after=10, first_line_indent=Pt(0))

add_page_break()

# ====================================================================
# ---------- 1. 과제 목표 ----------
# ====================================================================
add_section_heading("1. 과제 목표")
add_body(
    "Wolfenstein 3D 스타일의 1인칭/3인칭 미로 슈팅 게임을 제작한다. DirectX 12 의 제한된 "
    "파이프라인(입력 조립기 → 정점 셰이더 → 픽셀 셰이더 → 출력 병합기) 안에서, "
    "절차적으로 생성된 미로를 탐험하며 등장하는 적 AI 를 모두 처치하는 것을 목표로 한다. "
    "타이틀 화면, 맵 선택 화면, 서로 다른 두 개의 게임 맵을 모두 구현하고, "
    "1인칭/3인칭 시점 전환, 화면 십자선 조준점, 라이프 바, 잔여 적 표시 등 기본적인 "
    "슈팅 게임 HUD 를 함께 갖춘다."
)

# ====================================================================
# ---------- 2. 가정 ----------
# ====================================================================
add_section_heading("2. 가정")
for line in [
    "0) C++20 / Windows 10 이상 / Direct3D 12 사용. 한글 코멘트는 CP949 로 인코딩한다.",
    "1) DirectX 12 의 텍스처/지오메트리/테셀레이션 셰이더 등은 사용하지 않는다. 색상만 사용한다.",
    "2) 깊이 검사는 D3D12 깊이-스텐실 버퍼가 처리한다 (D24_UNORM_S8_UINT 포맷).",
    "3) 면 단위 평탄 음영(flat shading) 을 사용하여 각 면을 또렷이 구분한다.",
    "4) 라이팅은 단일 방향광 디퓨즈(Lambert) + 환경광. 그림자/스페큘러 없음.",
    "5) 미로는 매 빌드마다 동일한 시드를 사용해 깊이 우선 탐색(DFS) 기반으로 자동 생성한다.",
    "6) 적 AI 는 한 번 시야에 들어온 플레이어를 영구적으로 추격한다 (시뮬레이션 단순화).",
    "7) 게임은 타이틀 화면에서 시작하여 맵 선택 → 스테이지 진입 흐름을 따른다.",
    "8) HUD 는 NDC 좌표를 직접 사용하는 별도 셰이더로 그려, 카메라/맵 회전과 무관하게 화면에 고정된다.",
    "9) 마우스 입력은 화면 중앙 캡쳐(SetCursorPos) 방식으로 무한 회전을 구현한다.",
]:
    add_para(line, size=11, space_after=3, first_line_indent=Pt(0))

# ====================================================================
# ---------- 3. 조작법 ----------
# ====================================================================
add_section_heading("3. 조작법")

add_subsection_heading("1) 타이틀 화면")
add_body("화면 중앙의 \"GAME START\" 버튼을 좌 클릭하면 맵 선택 화면으로 진입한다.")

add_subsection_heading("2) 맵 선택 화면")
add_body(
    "회전하는 두 개의 미로 미니어처 중 하나를 좌 클릭하면 해당 맵의 게임 스테이지로 진입한다."
)

add_subsection_heading("3) 스테이지 화면")
for line in [
    "W / A / S / D : 전후/좌우 이동",
    "마우스 움직임 : 시점(Pitch/Yaw) 회전",
    "마우스 좌 클릭 : 소총 발사 (쿨다운 약 0.25초)",
    "스페이스 바 : 점프 (낮은 단차는 자동 점프)",
    "Tab 키 : 1인칭(FPS) ↔ 3인칭(TPS) 시점 전환",
    "ESC 키 : 마우스 캡처 해제 / 종료",
]:
    add_bullet(line)

# ====================================================================
# ---------- 4. 실행 결과 ----------
# ====================================================================
add_section_heading("4. 실행 결과")
add_body(
    "프로그램 실행 시 \"THE MAZE\" 타이틀 화면이 표시되고, 그 아래 회전하는 작은 미로 "
    "미리보기와 \"GAME START\" 버튼이 등장한다. 버튼 클릭 시 맵 선택 화면으로 전환되어 "
    "회전하는 두 개의 미니어처 미로가 나타나며, 이를 클릭하면 해당 맵의 1인칭 시점 게임이 "
    "시작된다. 게임 중에는 화면 중앙의 십자선, 좌상단의 잔여 적 카운트, 중앙 하단의 라이프 바, "
    "TPS 모드로 전환 시 어깨너머 시점의 플레이어 모델과 소총이 함께 표시된다. 모든 적을 "
    "처치하면 화면에 \"WIN\" 글자가 표시된 뒤 자동으로 타이틀로 복귀하고, 라이프가 0 이 되면 "
    "화면이 천천히 기울며 페이드 아웃되어 타이틀 화면으로 돌아간다."
)

add_page_break()

# ====================================================================
# ---------- 5. 구현 내용 (서술형) ----------
# ====================================================================
add_section_heading("5. 구현 내용")

# ---------- 1) 씬 구조와 흐름 ----------
add_subsection_heading("1) 씬 구조 분리와 게임 흐름 제어")
add_body(
    "가장 먼저 한 일은 게임 흐름을 표현할 씬 구조를 정리하는 것이었습니다. 1차 과제와 달리 "
    "이번 과제는 타이틀 / 맵 선택 / 게임 플레이 화면이 모두 있어야 하므로, CScene 안에 "
    "SceneState 라는 enum 을 두고 LANDING / MAP_SELECT / MAP1 / MAP2 네 가지 상태로 "
    "구분했습니다."
)
add_body(
    "각 씬은 자신만의 CObjectsShader 와 GameObject 목록을 m_vShaders 벡터의 슬롯에 보유하며, "
    "Render() 호출 시 현재 상태 슬롯만 그리는 방식입니다. 씬 사이 이동은 TransitionToScene("
    "newState) 한 줄로 처리하고, 게임 흐름은 \"GAME START 클릭 → MAP_SELECT → 미니어처 클릭 → "
    "MAP1/MAP2 → 사망 또는 모든 적 처치 → LANDING 복귀\" 의 사이클을 따릅니다."
)

# ---------- 2) 타이틀 텍스트와 시작 버튼 ----------
add_subsection_heading("2) 타이틀 화면의 3D 텍스트와 시작 버튼")
add_body(
    "타이틀 화면의 \"THE MAZE\" 글자와 \"GAME START\" 버튼은 모두 작은 큐브 묶음으로 표현되어 "
    "있습니다. 5×7 픽셀 격자에 알파벳 도안을 미리 등록한 LetterPatternTable 을 두고, 각 글자의 "
    "true 셀마다 작은 큐브 정점을 쌓아 글자 하나의 메시(CTextLetterMesh) 를 만듭니다. "
    "CTextStringObject 가 그 메시를 가로로 정렬해 단어가 되고, 각 글자는 sine 파형의 wave "
    "애니메이션 위상을 글자별로 오프셋해 부드러운 물결 효과를 만듭니다."
)
add_body(
    "버튼(CButtonObject) 은 텍스트 객체를 상속하여 글자 묶음의 바운딩 박스에 약간의 패딩을 더한 "
    "영역에 화면 좌표 클릭이 들어왔는지(HitTest) 검사합니다. 화면 좌표를 월드 좌표로 환산할 "
    "때는 클릭 위치의 NDC 좌표를 inverse(view × proj) 로 풀어 z = 0 평면과의 교점을 구하는 "
    "방식을 사용했습니다."
)

# ---------- 3) 미로 자동 생성 ----------
add_subsection_heading("3) 미로 자동 생성")
add_body(
    "\"서로 다른 두 개의 맵\" 요구사항을 만족하기 위해 미로는 매번 시드만 다른 절차적 생성 "
    "알고리즘으로 만들었습니다. 36×36 격자를 모두 벽('W') 으로 채운 뒤, 깊이 우선 탐색(DFS) "
    "으로 통로('.') 를 깎아 들어가며 한 줄짜리 미로 골격을 만듭니다. 그 다음 WidenCorridors() "
    "가 25% 확률로 인접한 벽을 추가로 허물어 통로의 폭에 변화를 주고, PlantRooms() 가 "
    "2×2 ~ 4×4 크기의 직사각형 방을 6~10개 무작위로 심어 미로에 \"방\" 의 개념을 만듭니다."
)
add_body(
    "골격 → 폭 확장 → 방 배치의 3단 절차 덕분에 한 줄 통로의 단조로움이 깨지고, 같은 알고리즘에 "
    "시드만 1, 2 로 바꿔주는 것만으로 MAP1, MAP2 가 시각적으로 전혀 다른 미로가 됩니다."
)

# ---------- 4) 단차/벽 높이 다양화 ----------
add_subsection_heading("4) 단차 지형과 벽 높이 다양화")
add_body(
    "요구사항 3 의 \"비-평면 지형\" 은 바닥과 벽 두 갈래로 나누어 처리했습니다."
)
add_body(
    "바닥의 단차는 PartitionFloorRegions() 에서 처리합니다. 보행 가능한 '.' 셀을 3~7 셀 크기의 "
    "작은 영역으로 무작위 그리디 BFS 로 분할한 뒤, 각 영역마다 0~3 단계의 단차 높이(한 단 = "
    "STEP_H = 0.7m)를 부여합니다. 인접 영역끼리는 가능하면 다른 높이를 선택하도록 그래프 컬러링 "
    "방식으로 보정해 시각적으로 평탄한 덩어리가 생기지 않게 했고, 스폰 셀 (1, 1) 이 속한 "
    "영역은 높이 0 으로 강제하여 시작 즉시 공중에 뜨는 현상을 방지했습니다."
)
add_body(
    "벽의 높이는 ComputeWallHeights() 에서 처리합니다. 동일한 'W' 셀로 이루어진 4방향 연결 "
    "영역(군집) 을 BFS 로 묶고, 군집의 크기가 10 셀 이상이면 8.0 / 9.4 / 10.8 / 12.2 중 하나의 "
    "통일된 높이를 부여합니다. \"같은 한 덩어리의 벽은 같은 높이\" 라는 규칙이 자연스럽고 "
    "도시 같은 외관을 만듭니다."
)

# ---------- 5) 캐릭터 이동과 충돌 ----------
add_subsection_heading("5) 캐릭터 이동과 충돌 처리")
add_body(
    "플레이어와 적은 모두 CCharacter 라는 중간 클래스를 공유합니다. 체력, 점프 물리, 소총 부착, "
    "XZ 분리 프로브 이동 같은 공통 기능을 이 클래스에 묶고, CPlayer 는 입력 처리와 반동 "
    "타이머만, CEnemyObject 는 AI 상태머신과 길찾기만 추가로 가집니다. 이 추상화 덕분에 "
    "적의 행동을 수정할 때 플레이어 코드를 건드릴 일이 거의 없습니다."
)
add_body(
    "실제 이동 처리는 X 축과 Z 축을 분리해 각각 IsBlockedInMap() 으로 검사하는 방식입니다. "
    "한 축이 막혀도 다른 축은 통과시켜 벽에 비스듬히 다가갈 때 자연스럽게 미끄러집니다. "
    "만약 막힌 셀의 윗면이 발 높이 + STEP_UP_TOLERANCE 보다 낮으면 \"올라설 수 있는 단차\" 로 "
    "판단해 통과시키고, 그보다 높으면 자동으로 점프를 시작합니다. 점프 초기 속도는 "
    "sqrt(2 g h) 로 계산해 정확히 단차 위에 도달할 수 있는 최소 속도를 부여하도록 했습니다."
)
add_code(
    "// XZ 분리 프로브 이동 — GameObject.cpp::TryMoveXZ\n"
    "if (dx != 0.0f) {\n"
    "    const float probeX = pos.x + dx + (dx > 0.0f ? fRadius : -fRadius);\n"
    "    if (!IsBlockedInMap(m_eSceneState, probeX, pos.z, kFeetY)) {\n"
    "        pos.x += dx;\n"
    "    } else {\n"
    "        const float nextFloor = GetFloorHeightAt(m_eSceneState, probeX, pos.z);\n"
    "        if (nextFloor > kFeetY + 0.05f) bBlockedByStep = true;\n"
    "    }\n"
    "}\n"
    "if (bBlockedByStep && !m_bJumping && m_fJumpCooldown <= 0.0f) {\n"
    "    const float jumpH = max(stepHeight - kFeetY + 0.3f, 0.5f);\n"
    "    m_fVerticalVelocity = sqrtf(2.0f * kGravity * jumpH);\n"
    "    m_bJumping = true;\n"
    "}"
)
add_body(
    "스페이스 키 점프와 단차 자동 점프는 모두 동일한 ApplyJumpPhysics() 메커니즘을 공유합니다. "
    "매 프레임 중력(g = 20.0) 을 누적해 수직 속도를 감소시키고, 바닥에 닿으면 m_bJumping = "
    "false 로 복귀시키는 식입니다."
)

# ---------- 6) 1인칭/3인칭 카메라 ----------
add_subsection_heading("6) 1인칭/3인칭 카메라와 마우스 룩")
add_body(
    "CCamera 클래스는 m_fPitch / m_fYaw 두 각도만 보유하고, 이 값으로부터 look / right / up "
    "정규직교 기저를 매 프레임 재구성합니다(RegenerateViewMatrix). FPS 모드에서는 카메라 "
    "위치가 플레이어 위치와 동일(눈높이 MAP_EYE_HEIGHT = 3.5)이지만, TPS 모드에서는 플레이어 "
    "뒤쪽으로 약 7m 떨어진 어깨너머 위치에 놓이며 SetPositionAndTarget 으로 플레이어를 "
    "정확히 바라보도록 갱신합니다."
)
add_body(
    "3인칭 카메라가 벽을 뚫고 들어가는 문제를 막기 위해 ClampDistanceAgainstWalls() 라는 "
    "스프링 암 함수를 만들었습니다. 카메라가 가야 할 거리만큼의 수평 광선을 0.25 타일 간격으로 "
    "마칭하면서 첫 벽 충돌점까지의 거리를 잘라 반환합니다. 카메라는 그 거리만큼만 뒤로 빠지므로 "
    "벽 안쪽으로 들어가지 않습니다."
)
add_body(
    "1인칭 마우스 룩은 m_bMouseCaptured 플래그가 true 일 때 매 프레임 마우스 좌표를 측정한 뒤 "
    "화면 중앙으로 SetCursorPos() 하는 방식입니다. 모니터 가장자리에 부딪혀 멈추는 문제 없이 "
    "무한 회전이 가능하며, \"화면 중앙 좌표 - 현재 좌표 = 델타 픽셀\" 을 카메라의 "
    "Rotate(-fPitchDelta, fYawDelta) 에 전달합니다. Y 부호는 반전해서 마우스를 위로 올리면 "
    "시야도 위로 향하게 했습니다."
)

# ---------- 7) 적 AI ----------
add_subsection_heading("7) 적 AI ─ 상태머신, 시야 판정, A* 길찾기")
add_body(
    "적은 Idle_Move(패트롤 이동) / Idle_Pause(정지) / Pursue(추적) 세 가지 상태를 가지는 "
    "상태머신입니다. 매 프레임 HasLineOfSight() 로 플레이어와의 시야를 검사하고, 한 번이라도 "
    "시야에 잡히면 m_bAware = true 가 되어 영구적으로 Pursue 상태로 전환됩니다. "
    "(사용자 결정: 영구 추적 ─ 게임 진행을 단순화하기 위해 한 번 발견된 적은 끝까지 따라오게 "
    "했습니다.)"
)
add_body(
    "Pursue 상태에서는 직선으로 플레이어를 추적하는 것이 아니라 FindPathAStar() 로 그리드 셀 "
    "단위 최단 경로를 탐색합니다. 'W' 셀만 불통으로 처리하고 단차('1'~'3') 는 통과 가능 취급 "
    "(TryMoveXZ 의 점프 메커니즘이 단차를 자동으로 넘어가기 때문)합니다. 0.5초마다 경로를 "
    "재계산하면서 다음 웨이포인트로 이동하고, 발사 쿨다운이 끝나면 총알을 발사한 뒤 약 0.3초간 "
    "정지하여 \"조준\" 연출을 보입니다."
)
add_code(
    "// A* 길찾기 핵심 ─ Maps.cpp::FindPathAStar\n"
    "auto isWalkable = [&](int c, int r) {\n"
    "    return (c>=0 && c<cols && r>=0 && r<rows) ? (grid[r][c] != 'W') : false;\n"
    "};\n"
    "auto heuristic = [&](int c, int r) {\n"
    "    return float(abs(c - goalC) + abs(r - goalR)) * TILE;  // 맨해튼\n"
    "};\n"
    "while (!openSet.empty() && nodesExpanded < nMaxNodes) {\n"
    "    int curIdx = openSet.top().second; openSet.pop();\n"
    "    if (curIdx == goalIdx) { /* 부모 추적으로 경로 복원 */ }\n"
    "    for (int i = 0; i < 4; ++i) {\n"
    "        float newG = gCost[curIdx] + TILE;\n"
    "        if (newG < gCost[nIdx]) {\n"
    "            gCost[nIdx] = newG; parent[nIdx] = curIdx;\n"
    "            openSet.emplace(newG + heuristic(nc, nr), nIdx);\n"
    "        }\n"
    "    }\n"
    "}"
)
add_body(
    "적의 스폰 위치는 PickEnemySpawnPositions() 가 결정합니다. 보행 가능한 셀 중 플레이어 시작 "
    "타일과의 Chebyshev 거리가 5 이상(약 20m) 인 것만 후보로 삼고, std::shuffle 로 섞은 뒤 "
    "최대 12 개를 뽑습니다. 즉시 사격당하지 않도록 시작 거리를 보장하는 방식입니다."
)

# ---------- 8) 사격과 충돌 ----------
add_subsection_heading("8) 사격 시스템과 일반화된 충돌")
add_body(
    "사격은 화면 중앙(십자선) 방향으로 광선을 쏴 첫 충돌점을 계산하는 방식으로 구현했습니다. "
    "GetAimTargetPoint() 가 우선 살아있는 적의 AABB 와 슬랩 라인 교차를 검사하고, 그 다음 벽 "
    "셀과 step-march 검사를 한 뒤 두 결과 중 가까운 것을 채택합니다. 총구는 소총 메시의 +Z 축 "
    "에서 1.25 단위 앞에 위치하며, 조준 대상점을 향해 정확히 발사됩니다. 발사 직후에는 "
    "m_fRecoilTimer 가 0 으로 리셋되어 3-키프레임 반동 애니메이션(뒤로 후진 → 절반 복귀 → "
    "원위치)이 시작됩니다."
)
add_body(
    "총알과 적, 적 총알과 플레이어의 충돌은 일반화된 충돌 시스템으로 처리합니다. 객체마다 "
    "EObjectTag enum 값(Bullet / Enemy / Player / EnemyBullet 등) 을 부여하고, "
    "Collision::CheckCollisions(objects, tagA, tagB, callback) 으로 두 태그 쌍의 모든 AABB "
    "겹침을 검사합니다. 충돌 페어로 등록하지 않은 쌍(예: Bullet vs EnemyBullet) 은 서로 "
    "통과하므로, 새 객체 타입이 추가될 때 새 태그 값과 한 줄의 콜백 등록만으로 확장 가능한 "
    "구조입니다."
)
add_code(
    "// 일반화된 충돌 검사 ─ Scene.cpp::AnimateObjects\n"
    "Collision::CheckCollisions(\n"
    "    m_vShaders[idx].GetObjects(),\n"
    "    EObjectTag::Bullet, EObjectTag::Enemy,\n"
    "    [](CGameObject* a, CGameObject* b) {\n"
    "        auto* pEnemy = static_cast<CEnemyObject*>(b);\n"
    "        if (pEnemy->IsDying()) return;\n"
    "        a->Kill();           // 총알 소멸\n"
    "        b->OnHit(a);         // 적 HP 감산 → 0이면 사망 애니메이션 진입\n"
    "    });"
)

# ---------- 9) 라이팅과 평탄 음영 ----------
add_subsection_heading("9) 디퓨즈 라이팅과 면 단위 평탄 음영")
add_body(
    "라이팅은 단일 방향광 디퓨즈(Lambert) + 환경광 모델로 구현했습니다. cbLightInfo(b2) 에 "
    "정규화된 광선 방향, 색상, 환경광 색상을 PS 전용 root constant 로 업로드하고, PS 에서 "
    "N · L 을 계산해 베이스 색상에 곱하는 단순한 방식입니다."
)
add_body(
    "면 단위 평탄 음영(flat shading) 을 위해서는 큐브 메시의 정점을 8 개 → 24 개로 분리하는 "
    "작업이 필요했습니다. 기본 8 정점 큐브는 한 정점이 인접한 3 면에서 공유되어 노멀이 평균되며 "
    "부드러운 음영이 나오는데, 면마다 독립된 4 정점을 두고 그 면의 노멀(±X / ±Y / ±Z) 을 "
    "부여하면 면 사이가 또렷이 끊어집니다. 노멀은 위치/색 정점 버퍼와 분리된 slot 1 의 별도 "
    "정점 버퍼로 바인딩했습니다(CNormalVertex)."
)

# ---------- 10) 메시 통합 최적화 ----------
add_subsection_heading("10) 정적 미로 메시 통합 (드로우 콜 최적화)")
add_body(
    "초기 구현에서 미로의 각 큐브(36 × 36 × 2 ≈ 2592 개) 를 개별 GameObject 로 그리는 방식은 "
    "매 프레임 수천 번의 draw call 을 발생시켜 CPU 바운드 병목을 만들었습니다. 회전이 없는 "
    "정적 큐브의 정점을 월드 좌표로 미리 변환해 단일 정점/인덱스 버퍼로 통합하고"
    "(CMergedCubeMesh), 한 번의 DrawIndexedInstanced 로 모든 벽/바닥을 그리도록 변경했습니다. "
    "결과적으로 1800+ draw call 이 1 개로 줄어 GPU 부하와 CPU 명령 전송이 모두 극적으로 "
    "감소했습니다. 동적 객체(적/총알) 는 회전이 들어가므로 그대로 개별 객체로 유지합니다."
)

# ---------- 11) .obj 로딩 ----------
add_subsection_heading("11) Wavefront .obj 파일 로딩")
add_body(
    "플레이어와 적이 들고 있는 소총은 외부 .obj 파일에서 로드한 M16 모델입니다. CObjMesh 가 "
    "v / vn / f / o 토큰을 직접 파싱하고, n-gon(>3) 페이스는 fan triangulation 으로 삼각형으로 "
    "분할합니다. vn 이 없는 페이스는 정점 위치로부터 외적으로 면 노멀을 자동 계산합니다."
)
add_body(
    "모델 좌표계를 엔진 좌표계(LH, +Z = forward) 로 변환하기 위해 호출자가 변환 행렬을 "
    "주입하면 정점/노멀에 미리 베이크하는 구조이며, o 그룹 이름을 이용해 \"Stock\" → 우드, "
    "\"Mag\" → 짙은 메탈 등 부품별 색상을 자동 분배합니다."
)

# ---------- 12) 사망 애니메이션 ----------
add_subsection_heading("12) 적/플레이어 사망 애니메이션")
add_body(
    "적이 HP 0 에 도달하면 즉시 Kill() 하지 않고 m_fDeathTimer 를 1.0초로 설정해 \"죽는 중\" "
    "상태로 진입합니다. 0.2초 정지 후 0.8초 동안 사망 시점의 forward 축을 중심으로 90° 회전 "
    "시켜 옆으로 쓰러뜨리고, 동시에 Y 좌표를 0.7m 만큼 lerp 로 하강시켜 회전으로 줄어든 수직 "
    "AABB 만큼 바닥에 자연스럽게 안착시킵니다."
)
add_body(
    "소총도 같이 쓰러져야 자연스럽기 때문에, 사망 시점의 소총 위치를 본체-로컬 좌표로 캐시"
    "(R_local = R_world × inverse(B_world)) 한 뒤 매 프레임 새 본체 행렬에 곱해(R_world(t) = "
    "R_local × B_world(t)) 강체 부착된 것처럼 따라가게 했습니다."
)
add_code(
    "// 사망 시점 소총의 본체-로컬 행렬 캐시 ─ GameObject.cpp::CEnemyObject::OnHit\n"
    "if (m_pRifle) {\n"
    "    XMMATRIX matRifleWorld = XMLoadFloat4x4(&m_pRifle->GetWorldMatrixRef());\n"
    "    XMMATRIX matBodyWorld  = XMLoadFloat4x4(&m_xmf4x4World);\n"
    "    XMVECTOR vDet;\n"
    "    XMMATRIX matBodyInv = XMMatrixInverse(&vDet, matBodyWorld);\n"
    "    XMStoreFloat4x4(&m_xmf4x4DeathRifleLocal,\n"
    "        XMMatrixMultiply(matRifleWorld, matBodyInv));\n"
    "    // 이후 매 프레임: R_world(t) = R_local * B_world(t) 로 복원\n"
    "}"
)
add_body(
    "플레이어 사망은 m_fPlayerDeathTimer 로 1.5초 동안 카메라가 천천히 앞으로 기울며 입력이 "
    "차단되고, 만료 시 LANDING 화면으로 자동 복귀합니다."
)

# ---------- 13) HUD ----------
add_subsection_heading("13) HUD 시스템과 피격 비네트")
add_body(
    "HUD 요소(십자선, 라이프 바, 적 잔여 카운트 점, WIN 글자, 피격 비네트) 는 모두 NDC 좌표를 "
    "직접 사용하는 CHudShader 위에서 동작합니다. 정점이 이미 NDC([-1, 1]) 좌표로 저장되어 있어 "
    "VSHud 셰이더가 월드/뷰/투영을 거치지 않고 그대로 출력하며, 깊이 테스트와 컬링을 모두 OFF "
    "로 두어 어떤 화면 위치든 항상 그려집니다. CCrosshairMesh / CLifeBarMesh / CHudQuadMesh "
    "가 모두 같은 셰이더를 공유합니다."
)
add_body(
    "피격 시 화면 외곽이 짧게 빨갛게 깜빡이는 비네트 효과는 별도의 CHitOverlayShader 가 알파 "
    "블렌딩이 켜진 PSO 로 풀스크린 쿼드 한 번을 그리고, PS 에서 화면 중앙으로부터의 거리에 "
    "따라 smoothstep 으로 외곽 가중치를 만드는 방식으로 구현했습니다."
)
add_code(
    "// HLSL ─ 화면 외곽 피격 비네트 ─ Shaders.hlsl::PSHitOverlay\n"
    "float4 PSHitOverlay(VS_OVERLAY_OUTPUT input) : SV_TARGET {\n"
    "    float r = length(input.ndc) * 0.7071f;\n"
    "    float mask = smoothstep(0.45f, 1.0f, r);\n"
    "    float a    = mask * gHitFlash;\n"
    "    if (a < 0.005f) discard;                  // 거의 안 보이면 일찍 종료\n"
    "    return float4(0.85f, 0.05f, 0.05f, a);\n"
    "}"
)

# ---------- 14) 미니어처/승리 처리 ----------
add_subsection_heading("14) 맵 선택 미니어처와 승리 처리")
add_body(
    "맵 선택 화면에서는 두 개의 게임 맵을 그대로 작은 크기(scale 0.25)로 회전시키며 보여줍니다. "
    "RenderInParent() 가 자식 객체의 월드 행렬에 부모 변환(축소 + X축 45° 기울임 + Y축 회전 + "
    "좌/우 이동) 을 곱해 동일한 미로 메시를 미니어처로 재활용합니다. 마우스 호버 시 hover "
    "scale 1.18 을 추가 적용해 선택 가능 영역을 시각적으로 강조했습니다."
)
add_body(
    "게임 진행 중 잔여 적이 3 마리 이하가 되면 Scene::SetEnemyMarkersVisible(true) 가 호출되어 "
    "살아있는 적의 머리 위에 노란 기둥 마커가 표시됩니다. 마커는 14m 길이로 만들어 8m 높이의 "
    "벽 너머로도 보이게 했습니다. 모든 적을 처치하면 m_fVictoryTimer 가 시작되고 화면 중앙에 "
    "\"WIN\" 글자가 CHudQuadMesh 세그먼트 조합(W = 4 stroke, I = 3 stroke, N = 3 stroke) 으로 "
    "표시된 뒤 타이머 만료 시 LANDING 으로 자동 복귀합니다."
)

add_page_break()

# ====================================================================
# ---------- 6. 미흡한 점 ----------
# ====================================================================
add_section_heading("6. 미흡한 점")
for line in [
    "1) 적의 총알 회피 행동이 없어 플레이어가 자리만 잡으면 적이 쉽게 제거된다.",
    "2) 적과 적 사이의 충돌 처리가 없어 같은 좁은 통로에 여러 적이 겹쳐 들어갈 수 있다.",
    "3) 게임 사운드(BGM, 발사음, 피격음) 가 없다.",
    "4) 텍스처를 사용할 수 없어 큐브 색상만으로 미로의 시각적 다양성을 표현해야 했다.",
    "5) A* 길찾기가 모든 적에 대해 0.5초마다 재계산되므로 적이 많아지면 CPU 부하가 누적될 수 있다.",
    "6) 1인칭 모드에서 자신의 소총만 보이고 손/팔이 렌더링되지 않는다.",
    "7) 사망 시 \"GAME OVER\" 같은 명시적 텍스트 표시 없이 페이드만 적용되어 처음 보면 종료된 줄 모를 수 있다.",
]:
    add_para(line, size=11, space_after=3, first_line_indent=Pt(0))

# ====================================================================
# 저장
# ====================================================================
doc.save(OUTPUT_PATH)
print(f"Generated: {OUTPUT_PATH}")
